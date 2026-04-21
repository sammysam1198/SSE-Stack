import io
import re
from datetime import datetime, timezone

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas


def slugify_filename(value: str) -> str:
    value = (value or "").strip().lower()
    value = re.sub(r"[^a-z0-9]+", "_", value)
    value = re.sub(r"_+", "_", value).strip("_")
    return value or "artist"


def build_contract_basename(artist_name: str, contract_type: str) -> str:
    artist_slug = slugify_filename(artist_name)
    contract_slug = slugify_filename(contract_type)
    date_str = datetime.now(timezone.utc).strftime("%Y%m%d")
    return f"{artist_slug}_{contract_slug}_contract_{date_str}"


def build_contract_object_keys(artist_name: str, contract_type: str):
    base = build_contract_basename(artist_name, contract_type)
    contract_type_slug = slugify_filename(contract_type)

    return {
        "docx": f"docs/contracts/{contract_type_slug}/unsigned/{base}.docx",
        "pdf": f"docs/contracts/{contract_type_slug}/unsigned/{base}.pdf",
        "signed": f"docs/contracts/{contract_type_slug}/signed/{base}_signed.pdf",
    }


def build_docx_bytes(title: str, body_text: str) -> bytes:
    document = Document()
    document.add_heading(title or "Contract", level=1)

    for paragraph in str(body_text or "").split("\n\n"):
        cleaned = paragraph.strip()
        if cleaned:
            document.add_paragraph(cleaned)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_pdf_bytes(title: str, body_text: str) -> bytes:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    width, height = LETTER

    x = 0.8 * inch
    y = height - 0.9 * inch

    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(x, y, title or "Contract")
    y -= 0.35 * inch

    pdf.setFont("Helvetica", 11)
    max_width = 90

    paragraphs = str(body_text or "").split("\n\n")

    for paragraph in paragraphs:
        words = paragraph.strip().split()
        if not words:
            y -= 0.18 * inch
            continue

        line = []
        for word in words:
            test_line = " ".join(line + [word])
            if pdf.stringWidth(test_line, "Helvetica", 11) <= 6.9 * inch:
                line.append(word)
            else:
                pdf.drawString(x, y, " ".join(line))
                y -= 0.2 * inch
                line = [word]

                if y < 0.9 * inch:
                    pdf.showPage()
                    pdf.setFont("Helvetica", 11)
                    y = height - 0.9 * inch

        if line:
            pdf.drawString(x, y, " ".join(line))
            y -= 0.26 * inch

        y -= 0.08 * inch

        if y < 0.9 * inch:
            pdf.showPage()
            pdf.setFont("Helvetica", 11)
            y = height - 0.9 * inch

    pdf.save()
    return buffer.getvalue()